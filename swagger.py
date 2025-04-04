import streamlit as st
import boto3
import json
import uuid
import docx
import PyPDF2
from langchain.memory import ConversationBufferMemory

# AWS Configuration
AWS_REGION = "us-east-1"
S3_BUCKET_NAME = "your-security-bucket"
CLAUDE_MODEL_ID = "anthropic.claude-3-5-sonnet-20240620-v1:0"
DYNAMODB_TABLE = "ChatHistory"

# AWS Clients
s3_client = boto3.client("s3", region_name=AWS_REGION)
bedrock_client = boto3.client("bedrock-runtime", region_name=AWS_REGION)
dynamodb = boto3.resource("dynamodb", region_name=AWS_REGION)
table = dynamodb.Table(DYNAMODB_TABLE)

# Unique Session ID
if "session_id" not in st.session_state:
    st.session_state.session_id = str(uuid.uuid4())

# Initialize Chat Memory
if "memory" not in st.session_state:
    st.session_state.memory = ConversationBufferMemory()

# Streamlit UI
st.set_page_config(page_title="AWS Bedrock Chatbot", page_icon="🤖", layout="wide")
st.title("AWS Bedrock Chatbot 🤖")

# Extract text from uploaded files
def extract_text_from_file(file):
    text = ""
    if file.name.endswith(".txt"):
        text = file.getvalue().decode("utf-8")
    elif file.name.endswith(".pdf"):
        pdf_reader = PyPDF2.PdfReader(file)
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
    elif file.name.endswith(".docx"):
        doc = docx.Document(file)
        for para in doc.paragraphs:
            text += para.text + "\n"
    return text.strip()

# Get Claude Response
def get_bedrock_response(prompt, chat_history):
    full_prompt = f"""
    Previous Conversation: {chat_history}
    
    User Query: {prompt}
    
    Claude:
    """
    
    payload = {
        "anthropic_version": "bedrock-2023-05-31",
        "messages": [{"role": "user", "content": full_prompt}],
        "max_tokens": 500
    }

    try:
        response = bedrock_client.invoke_model(
            modelId=CLAUDE_MODEL_ID,
            contentType="application/json",
            accept="application/json",
            body=json.dumps(payload)
        )
        response_body = json.loads(response["body"].read().decode("utf-8"))

        if "content" in response_body and isinstance(response_body["content"], list):
            return "\n".join(item["text"] for item in response_body["content"] if isinstance(item, dict) and "text" in item)
        return "Error: Unexpected response format from Claude."

    except Exception as e:
        return f"Error calling Claude 3.5: {e}"

# Store Chat in DynamoDB
def store_chat_in_dynamodb(user_input, bot_response):
    try:
        table.put_item(
            Item={
                "session_id": st.session_state.session_id,
                "timestamp": str(uuid.uuid4()),
                "user_input": user_input,
                "bot_response": bot_response
            }
        )
    except Exception as e:
        st.warning(f"Error storing chat history: {e}")

# Sidebar
with st.sidebar:
    if st.button("🗑️ Clear Chat"):
        st.session_state.memory.clear()
        st.session_state.messages = []
        st.rerun()

# Chat Input and File Upload (Side-by-Side)
col1, col2 = st.columns([8, 1])

with col1:
    user_input = st.text_input("Type your message here...", key="input", on_change=lambda: st.session_state.update({"submitted": True}))

with col2:
    uploaded_file = st.file_uploader("➕", type=["txt", "pdf", "docx"], key="upload", help="Upload a document for AI summary")

if uploaded_file:
    # Upload file to S3
    file_path = f"uploads/{uploaded_file.name}"
    s3_client.upload_fileobj(uploaded_file, S3_BUCKET_NAME, file_path)
    st.success(f"Uploaded to S3: {uploaded_file.name}")

    # Extract and summarize text
    extracted_text = extract_text_from_file(uploaded_file)
    summary = get_bedrock_response(f"Summarize this document:\n\n{extracted_text}", "")
    st.write(summary)

# Handle Chat Input
if st.session_state.get("submitted"):
    if user_input.strip():
        # Append user input to chat history
        st.session_state.messages.append(("user", user_input))

        # Retrieve past chat context
        past_context = st.session_state.memory.load_memory_variables({}).get("history", "")

        # Get AI response
        with st.spinner("Thinking..."):
            response = get_bedrock_response(user_input, past_context)

        # Save in memory
        st.session_state.memory.save_context({"input": user_input}, {"output": response})

        # Store in DynamoDB
        store_chat_in_dynamodb(user_input, response)

        # Append bot response to chat history
        st.session_state.messages.append(("bot", response))

    # Reset Input Field
    st.session_state["submitted"] = False
    st.experimental_rerun()

# Display Chat History
chat_container = st.container()
with chat_container:
    for role, content in st.session_state.messages:
        if role == "user":
            st.markdown(f'<div style="text-align: right; background: #dcf8c6; padding: 10px; margin: 5px; border-radius: 10px;">{content}</div>', unsafe_allow_html=True)
        else:
            st.markdown(f'<div style="text-align: left; background: #f1f0f0; padding: 10px; margin: 5px; border-radius: 10px;">{content}</div>', unsafe_allow_html=True)
            
import streamlit as st
import boto3
import json
import uuid
import docx
import PyPDF2
from langchain.memory import ConversationBufferMemory

# AWS Configuration
AWS_REGION = "us-east-1"
S3_BUCKET_NAME = "your-security-bucket"
CLAUDE_MODEL_ID = "anthropic.claude-3-5-sonnet-20240620-v1:0"
DYNAMODB_TABLE = "ChatHistory"

# AWS Clients
s3_client = boto3.client("s3", region_name=AWS_REGION)
bedrock_client = boto3.client("bedrock-runtime", region_name=AWS_REGION)
dynamodb = boto3.resource("dynamodb", region_name=AWS_REGION)
table = dynamodb.Table(DYNAMODB_TABLE)

# Unique Session ID
if "session_id" not in st.session_state:
    st.session_state.session_id = str(uuid.uuid4())

# Initialize Chat Memory
if "memory" not in st.session_state:
    st.session_state.memory = ConversationBufferMemory()
if "messages" not in st.session_state:
    st.session_state.messages = []

# Streamlit UI
st.set_page_config(page_title="AWS Bedrock Chatbot", page_icon="🤖", layout="wide")
st.title("AWS Bedrock Chatbot 🤖")

# Extract text from uploaded files
def extract_text_from_file(file):
    text = ""
    if file.name.endswith(".txt"):
        text = file.getvalue().decode("utf-8")
    elif file.name.endswith(".pdf"):
        pdf_reader = PyPDF2.PdfReader(file)
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
    elif file.name.endswith(".docx"):
        doc = docx.Document(file)
        for para in doc.paragraphs:
            text += para.text + "\n"
    return text.strip()

# Get Claude Response
def get_bedrock_response(prompt, chat_history):
    full_prompt = f"""
    Previous Conversation: {chat_history}
    
    User Query: {prompt}
    
    Claude:
    """
    
    payload = {
        "anthropic_version": "bedrock-2023-05-31",
        "messages": [{"role": "user", "content": full_prompt}],
        "max_tokens": 500
    }

    try:
        response = bedrock_client.invoke_model(
            modelId=CLAUDE_MODEL_ID,
            contentType="application/json",
            accept="application/json",
            body=json.dumps(payload)
        )
        response_body = json.loads(response["body"].read().decode("utf-8"))

        if "content" in response_body and isinstance(response_body["content"], list):
            return "\n".join(item["text"] for item in response_body["content"] if isinstance(item, dict) and "text" in item)
        return "Error: Unexpected response format from Claude."

    except Exception as e:
        return f"Error calling Claude 3.5: {e}"

# Store Chat in DynamoDB
def store_chat_in_dynamodb(user_input, bot_response):
    try:
        table.put_item(
            Item={
                "session_id": st.session_state.session_id,
                "timestamp": str(uuid.uuid4()),
                "user_input": user_input,
                "bot_response": bot_response
            }
        )
    except Exception as e:
        st.warning(f"Error storing chat history: {e}")

# Sidebar with "Clear Chat" and "New Chat"
with st.sidebar:
    if st.button("🗑️ Clear Chat"):
        st.session_state.memory.clear()
        st.session_state.messages = []
        st.rerun()
    
    if st.button("➕ New Chat"):
        st.session_state.session_id = str(uuid.uuid4())  # New session ID
        st.session_state.memory.clear()
        st.session_state.messages = []
        st.rerun()

# Chat Input and File Upload (Side-by-Side)
col1, col2 = st.columns([8, 1])

with col1:
    user_input = st.text_input("Type your message here...", key="input", on_change=lambda: st.session_state.update({"submitted": True}))

with col2:
    uploaded_file = st.file_uploader("📄", type=["txt", "pdf", "docx"], key="upload", help="Upload a document for AI analysis")

if uploaded_file:
    # Upload file to S3
    file_path = f"uploads/{uploaded_file.name}"
    s3_client.upload_fileobj(uploaded_file, S3_BUCKET_NAME, file_path)
    st.success(f"Uploaded to S3: {uploaded_file.name}")

    # Extract and summarize text
    extracted_text = extract_text_from_file(uploaded_file)
    summary = get_bedrock_response(f"Summarize this document:\n\n{extracted_text}", "")
    st.write(summary)

# Handle Chat Input
if st.session_state.get("submitted"):
    if user_input.strip():
        # Append user input to chat history
        st.session_state.messages.append(("user", user_input))

        # Retrieve past chat context
        past_context = st.session_state.memory.load_memory_variables({}).get("history", "")

        # Get AI response
        with st.spinner("Thinking..."):
            response = get_bedrock_response(user_input, past_context)

        # Save in memory
        st.session_state.memory.save_context({"input": user_input}, {"output": response})

        # Store in DynamoDB
        store_chat_in_dynamodb(user_input, response)

        # Append bot response to chat history
        st.session_state.messages.append(("bot", response))

    # Reset Input Field
    st.session_state["submitted"] = False
    st.experimental_rerun()

# Display Chat History
chat_container = st.container()
with chat_container:
    for role, content in st.session_state.messages:
        if role == "user":
            st.markdown(f'<div style="text-align: right; background: #dcf8c6; padding: 10px; margin: 5px; border-radius: 10px;">{content}</div>', unsafe_allow_html=True)
        else:
            st.markdown(f'<div style="text-align: left; background: #f1f0f0; padding: 10px; margin: 5px; border-radius: 10px;">{content}</div>', unsafe_allow_html=True)
            
import streamlit as st
import boto3
import json
import uuid
from langchain.memory import ConversationBufferMemory

# AWS Configuration
AWS_REGION = "us-east-1"
S3_BUCKET_NAME = "your-security-bucket"
CLAUDE_MODEL_ID = "anthropic.claude-3-5-sonnet-20240620-v1:0"
DYNAMODB_TABLE = "ChatHistory"

# AWS Clients
s3_client = boto3.client("s3", region_name=AWS_REGION)
bedrock_client = boto3.client("bedrock-runtime", region_name=AWS_REGION)
dynamodb = boto3.resource("dynamodb", region_name=AWS_REGION)
table = dynamodb.Table(DYNAMODB_TABLE)

# Unique Session ID
if "session_id" not in st.session_state:
    st.session_state.session_id = str(uuid.uuid4())

# Initialize Chat Memory
if "memory" not in st.session_state:
    st.session_state.memory = ConversationBufferMemory()

# Initialize Chat Messages
if "messages" not in st.session_state:
    st.session_state.messages = []

# Streamlit UI
st.set_page_config(page_title="AWS Bedrock Chatbot", page_icon="🤖", layout="wide")
st.title("AWS Bedrock Chatbot 🤖")

# Get Claude Response
def get_bedrock_response(prompt, chat_history, file_url=None):
    full_prompt = f"""
    Previous Conversation: {chat_history}
    
    User Query: {prompt}
    
    Claude:
    """
    
    if file_url:
        full_prompt += f"\n\n[Attached Document for Reference: {file_url}]"
    
    payload = {
        "anthropic_version": "bedrock-2023-05-31",
        "messages": [{"role": "user", "content": full_prompt}],
        "max_tokens": 500
    }

    try:
        response = bedrock_client.invoke_model(
            modelId=CLAUDE_MODEL_ID,
            contentType="application/json",
            accept="application/json",
            body=json.dumps(payload)
        )
        response_body = json.loads(response["body"].read().decode("utf-8"))

        if "content" in response_body and isinstance(response_body["content"], list):
            return "\n".join(item["text"] for item in response_body["content"] if isinstance(item, dict) and "text" in item)
        return "Error: Unexpected response format from Claude."

    except Exception as e:
        return f"Error calling Claude 3.5: {e}"

# Store Chat in DynamoDB
def store_chat_in_dynamodb(user_input, bot_response):
    try:
        table.put_item(
            Item={
                "session_id": st.session_state.session_id,
                "timestamp": str(uuid.uuid4()),
                "user_input": user_input,
                "bot_response": bot_response
            }
        )
    except Exception as e:
        st.warning(f"Error storing chat history: {e}")

# Sidebar with Chat Control Buttons
with st.sidebar:
    st.subheader("Chat Controls")
    
    # New Chat Button
    if st.button("🆕 New Chat"):
        st.session_state.memory.clear()
        st.session_state.messages = []
        st.session_state.session_id = str(uuid.uuid4())
        st.experimental_rerun()

    # Clear Chat Button
    if st.button("🗑️ Clear Chat"):
        st.session_state.memory.clear()
        st.session_state.messages = []
        st.experimental_rerun()

# Chat Input and File Upload (Side-by-Side)
col1, col2 = st.columns([8, 2])

with col1:
    user_input = st.text_input("Type your message here...", key="input", on_change=lambda: st.session_state.update({"submitted": True}))

with col2:
    uploaded_file = st.file_uploader("📎 Upload File", type=["txt", "pdf", "docx"], key="upload", help="Attach a document for Bedrock analysis")

# Handle File Upload
file_url = None
if uploaded_file:
    file_path = f"uploads/{uploaded_file.name}"
    s3_client.upload_fileobj(uploaded_file, S3_BUCKET_NAME, file_path)
    file_url = f"https://{S3_BUCKET_NAME}.s3.{AWS_REGION}.amazonaws.com/{file_path}"
    st.success(f"Uploaded to S3: [View File]({file_url})")

# Handle Chat Input
if st.session_state.get("submitted"):
    if user_input.strip():
        # Append user input to chat history
        st.session_state.messages.append(("user", user_input))

        # Retrieve past chat context
        past_context = st.session_state.memory.load_memory_variables({}).get("history", "")

        # Get AI response
        with st.spinner("Thinking..."):
            response = get_bedrock_response(user_input, past_context, file_url)

        # Save in memory
        st.session_state.memory.save_context({"input": user_input}, {"output": response})

        # Store in DynamoDB
        store_chat_in_dynamodb(user_input, response)

        # Append bot response to chat history
        st.session_state.messages.append(("bot", response))

    # Reset Input Field
    st.session_state["submitted"] = False
    st.experimental_rerun()

# Display Chat History
st.subheader("💬 Chat History")
chat_container = st.container()
with chat_container:
    for role, content in st.session_state.messages:
        if role == "user":
            st.markdown(f'<div style="text-align: right; background: #dcf8c6; padding: 10px; margin: 5px; border-radius: 10px;">{content}</div>', unsafe_allow_html=True)
        else:
            st.markdown(f'<div style="text-align: left; background: #f1f0f0; padding: 10px; margin: 5px; border-radius: 10px;">{content}</div>', unsafe_allow_html=True)
            
import streamlit as st
import boto3
import json
import uuid
from langchain.memory import ConversationBufferMemory

# AWS Configuration
AWS_REGION = "us-east-1"
CLAUDE_MODEL_ID = "anthropic.claude-3-5-sonnet-20240620-v1:0"
S3_BUCKET_NAME = "your-security-bucket"  # Update with your S3 bucket name
SECURITY_GUIDELINES_FILE = "security_guidelines.txt"  # Stored in S3

# Initialize AWS Clients
bedrock_client = boto3.client("bedrock-runtime", region_name=AWS_REGION)
s3_client = boto3.client("s3", region_name=AWS_REGION)

# Streamlit Page Config
st.set_page_config(page_title="AWS Bedrock Chatbot", page_icon="🤖", layout="wide")

# Handle Unique Chat Sessions via URL Parameters
query_params = st.query_params
chat_id = query_params.get("chat_id", None)

if chat_id is None:
    chat_id = str(uuid.uuid4())  # Generate new chat ID
    st.experimental_set_query_params(chat_id=chat_id)

# Initialize Memory for the Chat Session
if f"memory_{chat_id}" not in st.session_state:
    st.session_state[f"memory_{chat_id}"] = ConversationBufferMemory()
if f"messages_{chat_id}" not in st.session_state:
    st.session_state[f"messages_{chat_id}"] = []

# Function to Fetch Security Guidelines
def get_security_guidelines():
    try:
        response = s3_client.get_object(Bucket=S3_BUCKET_NAME, Key=SECURITY_GUIDELINES_FILE)
        return response["Body"].read().decode("utf-8")
    except Exception as e:
        return f"Error fetching security guidelines: {e}"

# Function to Call AWS Bedrock (Claude 3.5)
def get_bedrock_response(prompt, chat_history, security_guidelines):
    full_prompt = f"""
    Previous Conversation: {chat_history}
    
    User Query: {prompt}
    
    Security Guidelines Reference:
    {security_guidelines}

    Claude:
    """
    payload = {
        "anthropic_version": "bedrock-2023-05-31",
        "messages": [{"role": "user", "content": full_prompt}],
        "max_tokens": 500
    }
    try:
        response = bedrock_client.invoke_model(
            modelId=CLAUDE_MODEL_ID,
            contentType="application/json",
            accept="application/json",
            body=json.dumps(payload)
        )
        response_body = json.loads(response["body"].read().decode("utf-8"))
        return response_body.get("content", ["Error: No response"])[0]
    except Exception as e:
        return f"Error calling Claude 3.5: {e}"

# Sidebar Controls
with st.sidebar:
    st.header("🤖 Chatbot Settings")
    
    # New Chat Button - Opens a New Window with a New Chat ID
    if st.button("➕ New Chat"):
        new_chat_id = str(uuid.uuid4())
        st.experimental_set_query_params(chat_id=new_chat_id)
        st.rerun()

    # Clear Chat for the Current Session
    if st.button("🗑️ Clear Chat"):
        st.session_state[f"memory_{chat_id}"].clear()
        st.session_state[f"messages_{chat_id}"] = []
        st.rerun()

# Chat UI
st.title("AWS Bedrock Chatbot 🤖")
st.markdown(f"### 🆔 Chat ID: `{chat_id[:8]}` (You can switch back anytime)")

# Chat History Display
chat_container = st.container()
with chat_container:
    for role, content in st.session_state[f"messages_{chat_id}"]:
        if role == "user":
            st.markdown(f'<div style="background-color:#dcf8c6;padding:10px;border-radius:10px;margin:5px 0;">{content}</div>', unsafe_allow_html=True)
        else:
            st.markdown(f'<div style="background-color:#f1f0f0;padding:10px;border-radius:10px;margin:5px 0;">{content}</div>', unsafe_allow_html=True)

# User Input
user_input = st.text_input("Type your message here...", key="input", on_change=lambda: st.session_state.update({"submitted": True}))

if st.session_state.get("submitted"):
    if user_input.strip():
        # Append user input to chat history
        st.session_state[f"messages_{chat_id}"].append(("user", user_input))

        # Retrieve past chat context
        past_context = st.session_state[f"memory_{chat_id}"].load_memory_variables({}).get("history", "")

        # Fetch Security Guidelines
        security_guidelines = get_security_guidelines()

        # Get AI Response
        with st.spinner("Thinking..."):
            response = get_bedrock_response(user_input, past_context, security_guidelines)

        # Store in Memory
        st.session_state[f"memory_{chat_id}"].save_context({"input": user_input}, {"output": response})

        # Append Bot Response
        st.session_state[f"messages_{chat_id}"].append(("bot", response))

    # Reset Input
    st.session_state["submitted"] = False
    st.rerun()