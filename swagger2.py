st.markdown("""<style>div[data-testid="stFileUploadDropzone"] { visibility: hidden; height: 0; }</style>""", unsafe_allow_html=True)

# Display the custom upload button
st.markdown("<label class='upload-btn' style='display:inline-block;width:50px;height:50px;background:#0078ff;border-radius:50%;text-align:center;font-size:30px;color:white;line-height:50px;cursor:pointer;user-select:none;'>+</label>", unsafe_allow_html=True)

# Actual file uploader (hidden but functional)
uploaded_file = st.file_uploader("Upload file", type=["txt", "pdf", "docx"], key="file_uploader", label_visibility="hidden")

if uploaded_file:
    st.success("File uploaded successfully!")

st.markdown("""<style>div[data-testid="stFileUploadDropzone"] { visibility: hidden; height: 0; } .upload-btn { display: inline-block; width: 50px; height: 50px; background-color: #0078ff; border-radius: 50%; text-align: center; font-size: 30px; color: white; line-height: 50px; cursor: pointer; user-select: none; }</style><script>function triggerFileUpload() { document.querySelector('input[type=file]').click(); }</script><label class="upload-btn" onclick="triggerFileUpload()">+</label>""", unsafe_allow_html=True)

st.markdown("""<style>div[data-testid="stFileUploadDropzone"] { visibility: hidden; height: 0; } .upload-btn { display: inline-block; width: 50px; height: 50px; background-color: #0078ff; border-radius: 50%; text-align: center; font-size: 30px; color: white; line-height: 50px; cursor: pointer; user-select: none; }</style><label for='file-upload' class='upload-btn'>+</label>""", unsafe_allow_html=True)

st.markdown('<style>.file-uploader{visibility:hidden;height:0;}.upload-btn{display:inline-block;width:50px;height:50px;background-color:#0078ff;border-radius:50%;text-align:center;font-size:30px;color:white;line-height:50px;cursor:pointer;}</style>', unsafe_allow_html=True)

st.markdown(
    """
    <style>
        div[data-testid="stFileUploader"] > label {
            display: none;
        }
        div[data-testid="stFileUploader"] > div {
            display: flex;
            justify-content: center;
            align-items: center;
            width: 50px;
            height: 50px;
            background-color: #ddd;
            border-radius: 50%;
            font-size: 24px;
            cursor: pointer;
        }
        div[data-testid="stFileUploader"] > div::before {
            content: "+";
            font-weight: bold;
            color: black;
        }
    </style>
    """,
    unsafe_allow_html=True
)



import streamlit as st
import boto3
import json
import uuid
import PyPDF2
import docx
from langchain.memory import ConversationBufferMemory

# AWS Configuration
AWS_REGION = "us-east-1"
CLAUDE_MODEL_ID = "anthropic.claude-3-5-sonnet-20240620-v1:0"
S3_BUCKET_NAME = "your-security-bucket"
SECURITY_GUIDELINES_FILE = "security_guidelines.txt"

# Initialize AWS Clients
bedrock_client = boto3.client("bedrock-runtime", region_name=AWS_REGION)
s3_client = boto3.client("s3", region_name=AWS_REGION)

# Streamlit Page Config
st.set_page_config(page_title="AWS Bedrock Chatbot", page_icon="🤖", layout="wide")

# Handle Unique Chat Sessions
query_params = st.query_params
chat_id = query_params.get("chat_id", str(uuid.uuid4()))
st.query_params["chat_id"] = chat_id  

# Initialize Memory for the Chat Session
if f"memory_{chat_id}" not in st.session_state:
    st.session_state[f"memory_{chat_id}"] = ConversationBufferMemory()
if f"messages_{chat_id}" not in st.session_state:
    st.session_state[f"messages_{chat_id}"] = []
if f"uploaded_file_content_{chat_id}" not in st.session_state:
    st.session_state[f"uploaded_file_content_{chat_id}"] = ""

# Function to Extract Text from Uploaded Files
def extract_text_from_file(uploaded_file):
    file_extension = uploaded_file.name.split(".")[-1].lower()
    
    if file_extension == "txt":
        return uploaded_file.read().decode("utf-8")
    elif file_extension == "pdf":
        reader = PyPDF2.PdfReader(uploaded_file)
        return " ".join([page.extract_text() for page in reader.pages if page.extract_text()])
    elif file_extension == "docx":
        doc = docx.Document(uploaded_file)
        return " ".join([para.text for para in doc.paragraphs])
    else:
        return "Unsupported file type."

# Function to Fetch Security Guidelines from S3
def get_security_guidelines():
    try:
        response = s3_client.get_object(Bucket=S3_BUCKET_NAME, Key=SECURITY_GUIDELINES_FILE)
        return response["Body"].read().decode("utf-8")
    except Exception as e:
        return f"Error fetching security guidelines: {e}"

# Function to Call AWS Bedrock (Claude 3.5) with Security Guidelines & File Context
def get_bedrock_response(prompt, chat_history, file_content):
    # Fetch security guidelines from S3
    security_guidelines = get_security_guidelines()

    # Combine chat history, security guidelines, and file content
    full_prompt = f"""
    Previous Conversation: {chat_history}
    
    Security Guidelines:
    {security_guidelines}

    {f"Uploaded File Content:\n{file_content}" if file_content else ""}

    User Query: {prompt}
    
    Claude:
    """
    
    payload = {
        "anthropic_version": "bedrock-2023-05-31",
        "messages": [{"role": "user", "content": full_prompt}],
        "max_tokens": 500
    }

    # DEBUG: Show Final Prompt Sent to Claude
    st.markdown("### 📝 Claude's Final Prompt:")
    st.json(payload)

    try:
        response = bedrock_client.invoke_model(
            modelId=CLAUDE_MODEL_ID,
            contentType="application/json",
            accept="application/json",
            body=json.dumps(payload)
        )
        response_body = json.loads(response["body"].read().decode("utf-8"))

        if isinstance(response_body.get("content"), list):
            return " ".join([item.get("text", "") for item in response_body["content"] if isinstance(item, dict)])
        elif isinstance(response_body.get("content"), str):
            return response_body["content"]
        else:
            return "Error: Unexpected response format from Claude."
    except Exception as e:
        return f"Error calling Claude 3.5: {e}"

# Handle File Upload (Only Chat Input)
def handle_uploaded_file(uploaded_file, chat_id):
    if uploaded_file:
        extracted_text = extract_text_from_file(uploaded_file)
        if extracted_text:
            st.session_state[f"uploaded_file_content_{chat_id}"] = extracted_text
            st.success(f"📄 `{uploaded_file.name}` uploaded successfully!")

# **New Feature: Chat Input with File Upload Button**
col1, col2 = st.columns([8, 1])  

with col1:
    user_input = st.text_input("Type your message here...", key="input", on_change=lambda: st.session_state.update({"submitted": True}))

with col2:
    uploaded_file_chat = st.file_uploader("➕", type=["txt", "pdf", "docx"], label_visibility="collapsed")

# Process Chat Input File Upload
if uploaded_file_chat:
    handle_uploaded_file(uploaded_file_chat, chat_id)

# Fetch the latest file content before sending to Claude
file_content = st.session_state.get(f"uploaded_file_content_{chat_id}", "")

# Chat UI
st.title("AWS Bedrock Chatbot 🤖")
st.markdown(f"### 🆔 Chat ID: `{chat_id[:8]}` (Unique Session)")

# Display Chat History
chat_container = st.container()
with chat_container:
    for role, content in st.session_state[f"messages_{chat_id}"]:
        if role == "user":
            st.markdown(f'<div style="background-color:#dcf8c6;padding:10px;border-radius:10px;margin:5px 0;">{content}</div>', unsafe_allow_html=True)
        else:
            st.markdown(f'<div style="background-color:#f1f0f0;padding:10px;border-radius:10px;margin:5px 0;">{content}</div>', unsafe_allow_html=True)

# Process Chat Input
if st.session_state.get("submitted"):
    if user_input.strip():
        st.session_state[f"messages_{chat_id}"].append(("user", user_input))

        past_context = st.session_state[f"memory_{chat_id}"].load_memory_variables({}).get("history", "")

        # Get AI Response
        with st.spinner("Thinking..."):
            response = get_bedrock_response(user_input, past_context, file_content)

        # Store in Memory
        st.session_state[f"memory_{chat_id}"].save_context({"input": user_input}, {"output": response})

        # Append Bot Response
        st.session_state[f"messages_{chat_id}"].append(("bot", response))

    st.session_state["submitted"] = False
    st.rerun()

import gradio as gr
import boto3
import json
import uuid
import PyPDF2
import docx

# AWS Configuration
AWS_REGION = "us-east-1"
CLAUDE_MODEL_ID = "anthropic.claude-3-5-sonnet-20240620-v1:0"

# Initialize AWS Bedrock Client
bedrock_client = boto3.client("bedrock-runtime", region_name=AWS_REGION)

# Function to Extract Text from Uploaded Files
def extract_text_from_file(file):
    if file is None:
        return ""

    file_extension = file.name.split(".")[-1].lower()

    if file_extension == "txt":
        return file.read().decode("utf-8")
    elif file_extension == "pdf":
        reader = PyPDF2.PdfReader(file)
        return " ".join([page.extract_text() for page in reader.pages if page.extract_text()])
    elif file_extension == "docx":
        doc = docx.Document(file)
        return " ".join([para.text for para in doc.paragraphs])
    else:
        return "Unsupported file type."

# Function to Call AWS Bedrock (Claude 3.5) with Chat & File Context
def get_bedrock_response(chat_history, user_input, uploaded_file):
    file_content = extract_text_from_file(uploaded_file) if uploaded_file else ""

    full_prompt = f"""
    Previous Conversation: {chat_history}

    User Query: {user_input}

    Uploaded File Context:
    {file_content}

    Claude:
    """
    payload = {
        "anthropic_version": "bedrock-2023-05-31",
        "messages": [{"role": "user", "content": full_prompt}],
        "max_tokens": 500
    }
    
    try:
        response = bedrock_client.invoke_model(
            modelId=CLAUDE_MODEL_ID,
            contentType="application/json",
            accept="application/json",
            body=json.dumps(payload)
        )
        response_body = json.loads(response["body"].read().decode("utf-8"))
        
        if isinstance(response_body.get("content"), list):
            return " ".join([item.get("text", "") for item in response_body["content"] if isinstance(item, dict)])
        elif isinstance(response_body.get("content"), str):
            return response_body["content"]
        else:
            return "Error: Unexpected response format from Claude."
    except Exception as e:
        return f"Error calling Claude 3.5: {e}"

# Gradio Chatbot Interface
def chat_interface(chat_history, user_input, uploaded_file):
    response = get_bedrock_response(chat_history, user_input, uploaded_file)
    chat_history.append((user_input, response))
    return chat_history, None  # Clears input field after sending

# Gradio App Layout
with gr.Blocks() as app:
    gr.Markdown("### 🤖 AWS Bedrock Chatbot with Claude 3.5")

    chatbot = gr.Chatbot(label="AWS Bedrock Claude 3.5 Chatbot")
    with gr.Row():
        user_input = gr.Textbox(show_label=False, placeholder="Type your message here...", scale=8)
        file_upload = gr.File(label="📎", file_types=[".txt", ".pdf", ".docx"], scale=1)
        send_button = gr.Button("Send", scale=1)

    send_button.click(chat_interface, inputs=[chatbot, user_input, file_upload], outputs=[chatbot, user_input])
    user_input.submit(chat_interface, inputs=[chatbot, user_input, file_upload], outputs=[chatbot, user_input])

# Run the App
app.launch()